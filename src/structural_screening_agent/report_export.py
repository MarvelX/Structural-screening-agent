from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen.canvas import Canvas
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer
from typing import Iterable, Protocol

from structural_screening_agent.report_generator import ReportPreview


class _PreviewSectionLike(Protocol):
    heading: str
    items: list[str]


def _iter_sections_items(sections: Iterable[_PreviewSectionLike]):
    for section in sections:
        yield ("heading", section.heading)
        for item in section.items:
            if item.startswith("### "):
                yield ("subheading", item[4:])
            elif item.startswith("- "):
                yield ("bullet", item[2:])
            else:
                yield ("paragraph", item)


def _iter_preview_items(preview: ReportPreview):
    yield from _iter_sections_items(preview.sections)


def _is_heading(section_heading: str, zh_text: str, en_text: str) -> bool:
    return section_heading in {zh_text, en_text}


_EVIDENCE_MATRIX_HEADERS = [
    "发现项",
    "标题",
    "关联项",
    "证据类型",
    "来源",
    "位置",
    "摘录",
    "状态",
    "置信度",
]


def _is_evidence_matrix_section(section_heading: str) -> bool:
    return _is_heading(section_heading, "发现项证据矩阵", "Finding Evidence Matrix")


def _parse_evidence_matrix_item(item: str) -> dict[str, str]:
    values = {header: "" for header in _EVIDENCE_MATRIX_HEADERS}
    parts = [part.strip() for part in item.split(" | ") if part.strip()]
    if not parts:
        return values

    first_part = parts[0]
    if first_part.startswith("发现项 "):
        values["发现项"] = first_part.removeprefix("发现项 ").strip()
    else:
        values["发现项"] = first_part

    for part in parts[1:]:
        key, separator, value = part.partition(": ")
        if separator and key in values:
            values[key] = value.strip()
    return values


def _add_evidence_matrix_table(document: Document, items: list[str]) -> None:
    table = document.add_table(rows=1, cols=len(_EVIDENCE_MATRIX_HEADERS))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(_EVIDENCE_MATRIX_HEADERS):
        header_cells[index].text = header

    for item in items:
        values = _parse_evidence_matrix_item(item)
        row_cells = table.add_row().cells
        for index, header in enumerate(_EVIDENCE_MATRIX_HEADERS):
            row_cells[index].text = values.get(header, "")

    document.add_paragraph()


def _key_export_sections(preview: ReportPreview) -> tuple[list[tuple[str, str]], list]:
    conclusion_items: list[tuple[str, str]] = []
    calc_items: list[tuple[str, str]] = []
    remaining_sections = []

    for section in preview.sections:
        if _is_heading(section.heading, "初步结构结论", "Preliminary Structural Conclusion"):
            conclusion_items.append(("heading", section.heading))
            for item in section.items:
                conclusion_items.append(("bullet" if item.startswith("- ") else "paragraph", item[2:] if item.startswith("- ") else item))
            continue
        if _is_heading(section.heading, "简化计算结果", "Simplified Calculation Results"):
            calc_items.append(("heading", section.heading))
            kept = 0
            for item in section.items:
                normalized = item.strip()
                if "Formula" in normalized or "计算式" in normalized or "Result Unit" in normalized or "结果单位" in normalized:
                    continue
                calc_items.append(("bullet" if item.startswith("- ") else "paragraph", item[2:] if item.startswith("- ") else item))
                kept += 1
                if kept >= 6:
                    break
            continue
        remaining_sections.append(section)

    return [*conclusion_items, *calc_items], remaining_sections


def _cover_metadata_line(preview: ReportPreview) -> str | None:
    project_summary = next(
        (section for section in preview.sections if _is_heading(section.heading, "项目概况", "Project Summary")),
        None,
    )
    screening_inputs = next(
        (section for section in preview.sections if _is_heading(section.heading, "主案例筛查项", "Main-Case Screening Inputs")),
        None,
    )
    if project_summary is None and screening_inputs is None:
        return None

    fields: list[str] = []
    if project_summary is not None:
        for item in project_summary.items:
            if item.startswith("规范体系:") or item.startswith("Design Standard Context:"):
                fields.append(item)
            elif item.startswith("建筑类型:") or item.startswith("Building Type:"):
                fields.append(item)
    if screening_inputs is not None:
        for item in screening_inputs.items:
            if item.startswith("建筑跨度") or item.startswith("Building Span"):
                fields.append(item)
            elif item.startswith("柱距") or item.startswith("Column Spacing"):
                fields.append(item)
    if not fields:
        return None
    return " | ".join(fields[:4])


def build_docx_report_bytes(preview: ReportPreview) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal_style = document.styles["Normal"]
    normal_style.font.size = Pt(10.5)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(preview.title)
    footer_run.font.size = Pt(8.5)

    title = document.add_heading(preview.title, level=0)
    title.runs[0].font.size = Pt(18)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    metadata_line = _cover_metadata_line(preview)
    if metadata_line:
        metadata_paragraph = document.add_paragraph()
        metadata_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        metadata_paragraph.paragraph_format.space_after = Pt(8)
        metadata_paragraph.add_run(metadata_line)

    cover_items, remaining_sections = _key_export_sections(preview)
    for item_type, text in cover_items:
        if item_type == "heading":
            heading = document.add_heading(text, level=1)
            heading.runs[0].font.size = Pt(13)
        elif item_type == "bullet":
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.add_run(text)
        else:
            paragraph = document.add_paragraph(text)
            paragraph.paragraph_format.space_after = Pt(3)

    document.add_page_break()

    for section_item in remaining_sections:
        heading = document.add_heading(section_item.heading, level=1)
        heading.runs[0].font.size = Pt(13)
        if _is_evidence_matrix_section(section_item.heading):
            _add_evidence_matrix_table(document, section_item.items)
            continue
        for item_type, text in _iter_sections_items([section_item]):
            if item_type == "heading":
                continue
            if item_type == "subheading":
                subheading = document.add_heading(text, level=2)
                subheading.runs[0].font.size = Pt(11.5)
            elif item_type == "bullet":
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.paragraph_format.space_after = Pt(3)
                paragraph.add_run(text)
            else:
                paragraph = document.add_paragraph(text)
                paragraph.paragraph_format.space_after = Pt(3)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_pdf_report_bytes(preview: ReportPreview) -> bytes:
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ChineseTitle",
        parent=styles["Title"],
        fontName="STSong-Light",
        fontSize=18,
        leading=24,
    )
    heading_style = ParagraphStyle(
        "ChineseHeading",
        parent=styles["Heading2"],
        fontName="STSong-Light",
        fontSize=13.5,
        leading=19,
        spaceBefore=12,
        spaceAfter=6,
    )
    subheading_style = ParagraphStyle(
        "ChineseSubheading",
        parent=styles["Heading3"],
        fontName="STSong-Light",
        fontSize=11.5,
        leading=16,
        spaceBefore=8,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "ChineseBody",
        parent=styles["BodyText"],
        fontName="STSong-Light",
        fontSize=10.5,
        leading=15,
        spaceAfter=4,
    )
    bullet_style = ParagraphStyle(
        "ChineseBullet",
        parent=body_style,
        leftIndent=12,
        firstLineIndent=0,
        spaceAfter=3,
    )

    footer_style = ParagraphStyle(
        "ChineseFooter",
        parent=body_style,
        fontSize=8.5,
        leading=10,
        alignment=1,
    )

    output = BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        leftMargin=40,
        rightMargin=40,
        topMargin=40,
        bottomMargin=40,
    )

    story = [Paragraph(preview.title, title_style), Spacer(1, 8)]
    metadata_line = _cover_metadata_line(preview)
    if metadata_line:
        story.extend([Paragraph(metadata_line, footer_style), Spacer(1, 10)])
    cover_items, remaining_sections = _key_export_sections(preview)
    for item_type, text in cover_items:
        if item_type == "heading":
            story.append(Paragraph(text, heading_style))
        elif item_type == "bullet":
            story.append(Paragraph(text.replace("\n", "<br/>"), bullet_style))
        else:
            story.append(Paragraph(text.replace("\n", "<br/>"), body_style))
    story.extend([Spacer(1, 8), PageBreak()])

    for item_type, text in _iter_sections_items(remaining_sections):
        if item_type == "heading":
            story.append(Paragraph(text, heading_style))
        elif item_type == "subheading":
            story.append(Paragraph(text, subheading_style))
        elif item_type == "bullet":
            story.append(Paragraph(text.replace("\n", "<br/>"), bullet_style))
        else:
            story.append(Paragraph(text.replace("\n", "<br/>"), body_style))
    story.append(Spacer(1, 6))

    def draw_footer(canvas: Canvas, doc: SimpleDocTemplate) -> None:
        canvas.saveState()
        canvas.setFont("STSong-Light", 8.5)
        canvas.drawCentredString(A4[0] / 2, 22, preview.title)
        canvas.drawRightString(A4[0] - 40, 22, f"{canvas.getPageNumber()}")
        canvas.restoreState()

    document.build(story, onFirstPage=draw_footer, onLaterPages=draw_footer)
    return output.getvalue()
